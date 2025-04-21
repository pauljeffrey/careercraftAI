from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import os
import tempfile
from fpdf import FPDF
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
from typing import Optional, Union
import os
from reportlab.lib.pagesizes import A4
from schema import InterviewQuestions
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.enums import TA_LEFT


async def generate_resume_pdf(content: str, structure: dict = None) -> str:
    """
    Generate a PDF file from the optimized resume content,
    attempting to maintain the original structure if provided.
    """
    temp_dir = tempfile.gettempdir()
    output_path = os.path.join(temp_dir, "optimized_resume.pdf")
    
    # Create a PDF document
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Define styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='Heading1',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=12
    ))
    styles.add(ParagraphStyle(
        name='Normal_Justified',
        parent=styles['Normal'],
        alignment=4,  # 4 is for justified
        spaceBefore=6,
        spaceAfter=6
    ))
    
    # Create the content
    story = []
    
    if structure:
        # Try to maintain the original structure
        for section, lines in structure.items():
            # Add section header
            if section != "header":  # Assume "header" is not a visible section title
                p = Paragraph(section, styles['Heading1'])
                story.append(p)
                story.append(Spacer(1, 0.1 * inch))
            
            # Add section content
            for line in lines:
                p = Paragraph(line, styles['Normal_Justified'])
                story.append(p)
            
            story.append(Spacer(1, 0.2 * inch))
    else:
        # Just format the content as-is
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = Paragraph(para.replace('\n', '<br/>'), styles['Normal_Justified'])
                story.append(p)
                story.append(Spacer(1, 0.2 * inch))
    
    # Build the PDF
    doc.build(story)
    
    return output_path


def create_resume_docx(resume_data: dict, output_path: str, image_path: Optional[str] = None, image_bytes: Optional[bytes] = None):
    doc = Document()

    # Add image if provided
    if image_path:
        doc.add_picture(image_path, width=Inches(1.5))
    elif image_bytes:
        image = Image.open(BytesIO(image_bytes))
        temp_image_path = "temp_profile_image.png"
        image.save(temp_image_path)
        doc.add_picture(temp_image_path, width=Inches(1.5))
        os.remove(temp_image_path)

    doc.add_heading(resume_data['name'], level=1)
    doc.add_paragraph(f"{resume_data['email']} | {resume_data['phone']} | {resume_data['address']}")

    if resume_data.get('linkedin'):
        doc.add_paragraph(f"LinkedIn: {resume_data['linkedin']}")
    if resume_data.get('github'):
        doc.add_paragraph(f"GitHub: {resume_data['github']}")
    if resume_data.get('website'):
        doc.add_paragraph(f"Website: {resume_data['website']}")

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(resume_data['sections']['professional_summary'])

    def add_section(title, items, formatter):
        if items:
            doc.add_heading(title, level=2)
            for item in items:
                doc.add_paragraph(formatter(item), style='List Bullet')

    add_section("Skills", resume_data['sections']['skills'], lambda s: f"{s['name']} - {s.get('level', '')}")
    add_section("Work Experience", resume_data['sections']['experience'],
                lambda e: f"{e['title']} at {e['company']} ({e['start_date']} - {e['end_date']})\n{e['description']}\nAchievements: {e['achievements']}\nResponsibilities: {e['responsibilities']}")
    add_section("Education", resume_data['sections']['education'],
                lambda e: f"{e['degree']} - {e['institution']} ({e['start_date']} - {e['end_date']})\n{e['description']}")
    add_section("Certifications", resume_data['sections']['certifications'],
                lambda c: f"{c['name']} - {c['issuing_organization']} ({c['issue_date']})")
    add_section("Projects", resume_data['sections']['projects'],
                lambda p: f"{p['name']}: {p['description']} - Technologies: {', '.join(p['technologies_used'])}")
    add_section("Volunteer Experience", resume_data['sections']['volunteer_experience'],
                lambda v: f"{v['role']} at {v['organization']} ({v['start_date']} - {v['end_date']})\n{v['description']}")
    add_section("Publications", resume_data['sections']['publications'],
                lambda p: f"{p['title']} ({p['publication_date']}) - {p.get('link', '')}")
    add_section("Awards", resume_data['sections']['awards'],
                lambda a: f"{a['title']} - {a['organization']} ({a['date_received']})\n{a['description']}")
    add_section("References", resume_data['sections']['references'],
                lambda r: f"{r['name']}, {r['title']} at {r['organization']} - Contact: {r['contact_info']}")
    add_section("Additional Info", resume_data['sections']['additional_sections'],
                lambda i: f"{i['title']}: {i['description']}")

    if resume_data['sections']['languages']:
        doc.add_heading("Languages", level=2)
        doc.add_paragraph(", ".join(resume_data['sections']['languages']))

    if resume_data['sections']['interests']:
        doc.add_heading("Interests", level=2)
        doc.add_paragraph(", ".join(resume_data['sections']['interests']))

    doc.save(output_path)
    return 
    
    
def create_cover_letter_docx(output_path: str, cover_letter_data: dict):
    doc = Document()

    doc.add_paragraph(cover_letter_data['candidate_name'])
    doc.add_paragraph(cover_letter_data['candidate_address'])
    doc.add_paragraph("")

    doc.add_paragraph(cover_letter_data['recipient_name_or_office'])
    doc.add_paragraph(f"{cover_letter_data['recipient_title']}, {cover_letter_data['recipient_company']}")
    doc.add_paragraph(cover_letter_data['company_address'])
    doc.add_paragraph("")

    doc.add_heading(cover_letter_data['title'], level=1)
    doc.add_paragraph(cover_letter_data['body'])
    doc.add_paragraph("")
    doc.add_paragraph(cover_letter_data['closing'])

    doc.save(output_path)


def create_interview_question_pdf(questions: InterviewQuestions, output_path: str):
    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=40, leftMargin=40,
                            topMargin=50, bottomMargin=50)

    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    heading_style = styles['Heading2']
    heading_style.spaceAfter = 10
    normal_style.spaceAfter = 10

    content = []

    # Title
    content.append(Paragraph("📝 Interview Questions & Suggested Answers", styles['Title']))
    content.append(Spacer(1, 0.2 * inch))

    for idx, q in enumerate(questions.questions, start=1):
        question_type = f"👤 {q.Questioner} Question"
        content.append(Paragraph(f"<b>{idx}. {question_type}</b>", heading_style))
        content.append(Paragraph(f"<b>Question:</b> {q.question}", normal_style))

        if q.questionreason:
            content.append(Paragraph(f"<b>Why this question?</b> {q.questionreason}", normal_style))

        if q.suggestedAnswer:
            content.append(Paragraph(f"<b>Suggested Answer:</b> {q.suggestedAnswer}", normal_style))

        content.append(Spacer(1, 0.2 * inch))

    doc.build(content)
    doc.save(output_path)
    return 

def create_resume_pdf(resume_data: dict, output_path: str, image_path: Optional[str] = None, image_bytes: Optional[bytes] = None):
    doc = SimpleDocTemplate(output_path, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    content = []

    heading_style = ParagraphStyle('Heading', fontSize=14, spaceAfter=8, leading=16, alignment=TA_LEFT, bold=True)
    subheading_style = ParagraphStyle('SubHeading', fontSize=12, spaceAfter=6, leading=14, alignment=TA_LEFT, bold=True)
    body_style = ParagraphStyle('Body', fontSize=10, leading=13)

    # Add Image if present
    if image_path:
        content.append(RLImage(image_path, width=1.5 * inch, height=1.5 * inch))
        content.append(Spacer(1, 10))
    elif image_bytes:
        image = Image.open(BytesIO(image_bytes))
        temp_image_path = "temp_img.png"
        image.save(temp_image_path)
        content.append(RLImage(temp_image_path, width=1.5 * inch, height=1.5 * inch))
        content.append(Spacer(1, 10))
        os.remove(temp_image_path)

    # Personal Info
    content.append(Paragraph(f"<b>{resume_data['name']}</b>", heading_style))
    contact_info = f"{resume_data['email']} | {resume_data['phone']} | {resume_data['address']}"
    content.append(Paragraph(contact_info, body_style))
    for key in ['linkedin', 'github', 'website']:
        if resume_data.get(key):
            content.append(Paragraph(f"{key.capitalize()}: {resume_data[key]}", body_style))
    content.append(Spacer(1, 12))

    # Summary
    content.append(Paragraph("Professional Summary", subheading_style))
    content.append(Paragraph(resume_data['sections']['professional_summary'], body_style))
    content.append(Spacer(1, 12))

    def add_bullet_section(title, items, formatter):
        if items:
            content.append(Paragraph(title, subheading_style))
            for item in items:
                content.append(Paragraph(formatter(item), body_style))
                content.append(Spacer(1, 5))
            content.append(Spacer(1, 10))

    add_bullet_section("Skills", resume_data['sections']['skills'], lambda s: f"• {s['name']} - {s.get('level', '')}")
    add_bullet_section("Work Experience", resume_data['sections']['experience'],
                       lambda e: f"<b>{e['title']}</b> at {e['company']} ({e['start_date']} - {e['end_date']})<br/>{e['description']}<br/><i>Achievements:</i> {e['achievements']}<br/><i>Responsibilities:</i> {e['responsibilities']}")
    add_bullet_section("Education", resume_data['sections']['education'],
                       lambda e: f"<b>{e['degree']}</b> - {e['institution']} ({e['start_date']} - {e['end_date']})<br/>{e['description']}")
    add_bullet_section("Certifications", resume_data['sections']['certifications'],
                       lambda c: f"{c['name']} - {c['issuing_organization']} ({c['issue_date']})")
    add_bullet_section("Projects", resume_data['sections']['projects'],
                       lambda p: f"<b>{p['name']}</b>: {p['description']}<br/>Technologies: {', '.join(p['technologies_used'])}")
    add_bullet_section("Volunteer Experience", resume_data['sections']['volunteer_experience'],
                       lambda v: f"{v['role']} at {v['organization']} ({v['start_date']} - {v['end_date']})<br/>{v['description']}")
    add_bullet_section("Publications", resume_data['sections']['publications'],
                       lambda p: f"{p['title']} ({p['publication_date']}) - {p.get('link', '')}")
    add_bullet_section("Awards", resume_data['sections']['awards'],
                       lambda a: f"{a['title']} - {a['organization']} ({a['date_received']})<br/>{a['description']}")
    add_bullet_section("References", resume_data['sections']['references'],
                       lambda r: f"{r['name']}, {r['title']} at {r['organization']} - Contact: {r['contact_info']}")
    add_bullet_section("Additional Info", resume_data['sections']['additional_sections'],
                       lambda i: f"{i['title']}: {i['description']}")

    if resume_data['sections']['languages']:
        content.append(Paragraph("Languages", subheading_style))
        content.append(Paragraph(", ".join(resume_data['sections']['languages']), body_style))
        content.append(Spacer(1, 10))

    if resume_data['sections']['interests']:
        content.append(Paragraph("Interests", subheading_style))
        content.append(Paragraph(", ".join(resume_data['sections']['interests']), body_style))
        content.append(Spacer(1, 10))

    # Build PDF
    doc.build(content)
    return
